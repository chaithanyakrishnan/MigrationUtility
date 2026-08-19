"""
app/services/schema/extractor.py
P1.1 — Schema Extractor

Parses uploaded Relius/Frp schema files into a structured JSON format
that feeds the rest of the pipeline (P1.2 profiler, A1 RAG, P2.1 mapper).

Supported formats: SQL DDL, JSON, XLSX, PDF, DOCX, CSV, TXT, COBOL copybook (Frp only)
"""
from __future__ import annotations

import io
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import structlog


# PARSE-TIME WHITELIST — loaded from reference_data/relius_schema.json at startup.
# Purpose: prevents description words from being misidentified as table names
#          when parsing the Relius "Database Layout (By Tables)" PDF format.
# To update: edit app/reference_data/relius_schema.json — no code changes needed.
# To remove: pass valid_tables=None to _parse_relius_layout() and delete this block.
def _get_valid_tables() -> frozenset:
    """Load known Relius table names from the reference JSON (cached)."""
    try:
        from app.services.schema.profiler import get_valid_tables
        return get_valid_tables()
    except Exception:
        return frozenset()   # fallback: no whitelist (parser uses heuristics only)


logger = structlog.get_logger(__name__)


@dataclass
class ParsedField:
    table_name: str
    field_name: str
    data_type: str = ""
    nullable: bool = True
    is_pk: bool = False
    is_fk: bool = False
    references: Optional[str] = None  # "OTHER_TABLE.OTHER_FIELD"
    description: str = ""
    raw_line: str = ""


@dataclass
class ParsedTable:
    name: str
    fields: list[ParsedField] = field(default_factory=list)
    description: str = ""
    row_estimate: Optional[int] = None


@dataclass
class SchemaParseResult:
    tables: list[ParsedTable] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    raw_format: str = "unknown"

    @property
    def table_count(self) -> int:
        return len(self.tables)

    @property
    def field_count(self) -> int:
        return sum(len(t.fields) for t in self.tables)

    @property
    def fk_count(self) -> int:
        return sum(
            sum(1 for f in t.fields if f.is_fk)
            for t in self.tables
        )

    def to_dict(self) -> dict:
        return {
            "table_count": self.table_count,
            "field_count": self.field_count,
            "fk_count": self.fk_count,
            "format": self.raw_format,
            "tables": [
                {
                    "name": t.name,
                    "description": t.description,
                    "fields": [
                        {
                            "field": f.field_name,
                            "type": f.data_type,
                            "nullable": f.nullable,
                            "is_pk": f.is_pk,
                            "is_fk": f.is_fk,
                            "references": f.references,
                            "description": f.description,
                        }
                        for f in t.fields
                    ],
                }
                for t in self.tables
            ],
            "warnings": self.warnings,
        }


def _normalise_pypdf_va(line: str) -> str:
    """
    Rejoin 'VA' ligature splits introduced by pypdf on this PDF's font.
    pypdf renders the 'VA' ligature as 'V ' + rest, e.g.:
      COMBOVALUESDET  -> COMBOV ALUESDET
      CMSACTVASSOC    -> CMSACTV ASSOC
      VALRULENUM      -> V ALRULENUM

    Only applied to the TAIL (after the row number), never to description
    text. Uses a negative lookbehind to avoid matching digits inside type
    keywords like ARCHAR2 or NUMBER2.
    """
    import re as _re
    # Row number = digits NOT preceded by an uppercase letter (avoids "ARCHAR2")
    m = _re.search(r'(?<![A-Z])\d+[A-Z]', line)
    if m:
        prefix = line[:m.start()]
        tail   = line[m.start():]
        return prefix + _re.sub(r'V (A[A-Z0-9_]*)', lambda x: 'V' + x.group(1), tail)
    m2 = _re.search(r'(?<![A-Z])\d+\s+[A-Z]', line)
    if m2:
        prefix = line[:m2.start()]
        tail   = line[m2.start():]
        return prefix + _re.sub(r'V (A[A-Z0-9_]*)', lambda x: 'V' + x.group(1), tail)
    if _re.match(r'^\d+\s*[A-Z]', line):
        return _re.sub(r'V (A[A-Z0-9_]*)', lambda x: 'V' + x.group(1), line)
    return line


def _try_tail_with_split_fix(rest: str, valid_tables) -> tuple:
    """
    Recovery parser for lines where pypdf split a table or field name.
    pypdf splits uppercase ligatures: VALRULENUM → V ALRULENUM,
    CMSACTVASSOC → CMSA CTVASSOC etc.
    
    Strategy:
    1. Fix single-letter split in field name (V ALFIELD → VALFIELD)
    2. Try joining consecutive tokens to form a whitelisted table name,
       then join remaining tokens as the field name.
    
    Returns (table_name, field_name, length) — any may be empty string.
    """
    TAIL_PAT = re.compile(
        r"(\d+)([A-Z][A-Z0-9_]*(?:_[A-Z0-9]+)?)?\s+([A-Z_$#][A-Z0-9_$#]{0,49})"
        r"(?:\s+([\d(),]+))?\s*$"
    )

    # Fix 1: single-letter split in field name only
    # "1TABLE V ALFIELD 38" → "1TABLE VALFIELD 38"
    fixed = re.sub(
        r'(\d+[A-Z][A-Z0-9_]*\s+)([A-Z])\s+([A-Z][A-Z0-9_]*)',
        lambda m: m.group(1) + m.group(2) + m.group(3),
        rest
    )
    if fixed != rest:
        m = TAIL_PAT.search(fixed)
        if m:
            return (m.group(2) or "").upper() or "", m.group(3).upper(), m.group(4) or ""

    # Fix 2: split table name — try joining tokens to match whitelist
    rownum_m = re.search(r'(\d+)([A-Z].*)$', rest)
    if not rownum_m:
        return "", "", ""

    after_seq = rownum_m.group(2).strip()
    tokens = after_seq.split()

    for join_count in range(2, min(6, len(tokens))):
        candidate_table = "".join(tokens[:join_count]).upper()
        vt = valid_tables or set()
        if candidate_table in vt and len(tokens) > join_count:
            remaining = tokens[join_count:]
            length = ""
            if remaining and re.match(r'^[\d(),]+$', remaining[-1]):
                length = remaining[-1]
                remaining = remaining[:-1]
            # Join all remaining as field name (handles split field too)
            field_name = "".join(remaining).upper()
            if field_name:
                return candidate_table, field_name, length

    return "", "", ""


class SchemaExtractor:
    """
    Dispatches to format-specific parsers.
    Returns a SchemaParseResult regardless of input format.
    """

    def extract(self, filename: str, content: bytes) -> SchemaParseResult:
        ext = Path(filename).suffix.lower()
        logger.info("schema.extract.start", filename=filename, ext=ext, size=len(content))

        try:
            if ext in (".sql", ".ddl"):
                result = self._parse_sql(content.decode("utf-8", errors="replace"))
            elif ext == ".json":
                result = self._parse_json(content.decode("utf-8", errors="replace"))
            elif ext in (".xlsx", ".xls"):
                result = self._parse_xlsx(content)
            elif ext == ".csv":
                result = self._parse_csv(content.decode("utf-8", errors="replace"))
            elif ext in (".txt",):
                result = self._parse_txt(content.decode("utf-8", errors="replace"))
            elif ext == ".pdf":
                result = self._parse_pdf(content)
            elif ext == ".docx":
                result = self._parse_docx(content)
            elif ext in (".cbl", ".cob", ".cpy", ".cobol", ".copybook", ".cpb"):
                result = self._parse_cobol(content.decode("utf-8", errors="replace"))
            else:
                result = SchemaParseResult(warnings=[f"Unknown format '{ext}' — attempted text parse"])
                result = self._parse_txt(content.decode("utf-8", errors="replace"))

            result.raw_format = ext.lstrip(".")
            logger.info(
                "schema.extract.complete",
                filename=filename,
                tables=result.table_count,
                fields=result.field_count,
            )
            return result

        except Exception as exc:
            logger.error("schema.extract.error", filename=filename, error=str(exc))
            r = SchemaParseResult(warnings=[f"Parse error: {exc}"])
            r.raw_format = ext.lstrip(".")
            return r

    def extract_text(self, filename: str, content: bytes) -> tuple[str, str]:
        """
        Return best-effort *raw text* for a document, with OCR fallback for
        image-only PDFs, DOCX screenshots and standalone image files. Unlike
        `extract()` (which returns a schema structure), this is for consumers
        that parse their own layout — e.g. the transaction-card parser.
        Returns (text, warning).
        """
        ext = Path(filename).suffix.lower()
        try:
            if ext == ".pdf":
                warn = ""
                pages_text: list[str] = []
                try:
                    import pypdf
                    reader = pypdf.PdfReader(io.BytesIO(content))
                    pages_text = [p.extract_text() or "" for p in reader.pages]
                except Exception as e:
                    warn = f"PDF text error: {e}"
                text = "\n".join(pages_text)
                # OCR when the embedded text is thin — a text-layer heading alone
                # (e.g. just the report title) can exceed a tiny char threshold
                # while the actual body lives in page images. Detect that via
                # near-empty pages, then prefer whichever source recovers more.
                sparse_pages = sum(1 for t in pages_text if len(t.strip()) < 40)
                if len(text.strip()) < 200 or sparse_pages:
                    ocr, ocr_warn = self._ocr_pdf_pages(content)
                    if len(ocr.strip()) > len(text.strip()):
                        text = ocr
                    elif not text.strip():
                        warn = warn or ocr_warn
                return text, warn
            if ext == ".docx":
                warn = ""
                parts: list[str] = []
                try:
                    from docx import Document
                    doc = Document(io.BytesIO(content))
                    parts.extend(p.text for p in doc.paragraphs)
                    for tbl in doc.tables:
                        for row in tbl.rows:
                            parts.append("\t".join(c.text for c in row.cells))
                except Exception as e:
                    warn = f"DOCX text error: {e}"
                text_only = "\n".join(p for p in parts if p.strip())
                ocr, ocr_warn = self._ocr_docx_images(content)
                return "\n".join(t for t in (text_only, ocr) if t.strip()), (warn or ocr_warn)
            if ext in (".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".gif"):
                return self._ocr_image(content)
            return content.decode("utf-8", errors="replace"), ""
        except Exception as e:
            return "", f"Text extraction error: {e}"

    def _ocr_image(self, content: bytes) -> tuple[str, str]:
        """OCR a standalone image file (uploaded screenshot). (text, warning)."""
        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            return "", ("Image upload but OCR is not available — install pytesseract + "
                        "Pillow and the tesseract binary (`brew install tesseract`).")
        try:
            im = Image.open(io.BytesIO(content)).convert("L")
            w, h = im.size
            if w and w < 1600:
                scale = min(3.0, 1600 / w)
                im = im.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
            return pytesseract.image_to_string(im, config="--psm 6"), ""
        except Exception as e:
            return "", f"Image OCR error: {e}"

    # ── SQL DDL parser ─────────────────────────────────────────
    def _parse_sql(self, ddl: str) -> SchemaParseResult:
        """
        Parses Oracle/SQL Server DDL.
        Handles: CREATE TABLE, column definitions, PRIMARY KEY, FOREIGN KEY, REFERENCES.
        """
        result = SchemaParseResult()
        current_table: Optional[ParsedTable] = None

        # Normalise: strip comments, uppercase keywords
        lines = []
        for line in ddl.splitlines():
            stripped = line.strip()
            if stripped.startswith("--") or stripped.startswith("/*"):
                continue
            lines.append(stripped)

        text = " ".join(lines)
        # Split on CREATE TABLE statements
        table_blocks = re.split(r'CREATE\s+TABLE\s+', text, flags=re.IGNORECASE)

        for block in table_blocks[1:]:  # skip first empty split
            # Extract table name
            m = re.match(r'["[]?([\w.]+)["]]?\s*\(', block)
            if not m:
                continue
            table_name = m.group(1).upper().split(".")[-1]  # strip schema prefix
            table = ParsedTable(name=table_name)

            # Extract column definitions up to the closing )
            body_match = re.match(r'[\w."[\]]+\s*\((.+)', block, re.DOTALL)
            if not body_match:
                result.tables.append(table)
                continue

            body = body_match.group(1)
            # Find balanced closing paren
            depth = 1
            end = 0
            for i, ch in enumerate(body):
                if ch == "(": depth += 1
                elif ch == ")":
                    depth -= 1
                    if depth == 0:
                        end = i
                        break
            body = body[:end]

            # Parse column definitions
            for col_def in re.split(r",\s*(?=\w)", body):
                col_def = col_def.strip()
                if not col_def:
                    continue
                # Skip constraint lines
                if re.match(r'(CONSTRAINT|PRIMARY\s+KEY|FOREIGN\s+KEY|UNIQUE|CHECK|INDEX)', col_def, re.I):
                    # Extract FK references
                    fk_match = re.search(
                        r'FOREIGN\s+KEY\s*\((\w+)\)\s*REFERENCES\s+(\w+)\s*\((\w+)\)',
                        col_def, re.I
                    )
                    if fk_match:
                        fk_col, ref_table, ref_col = fk_match.groups()
                        for f in table.fields:
                            if f.field_name.upper() == fk_col.upper():
                                f.is_fk = True
                                f.references = f"{ref_table.upper()}.{ref_col.upper()}"
                    continue

                # Parse: field_name data_type [(size)] [NOT NULL] [DEFAULT ...]
                col_match = re.match(
                    r'["[]?(\w+)["]]?\s+([\w()., ]+?)(?:\s+(?:NOT\s+NULL|NULL|DEFAULT|PRIMARY|UNIQUE).*)?$',
                    col_def, re.I
                )
                if col_match:
                    fname, ftype = col_match.group(1), col_match.group(2).strip()
                    is_pk = "PRIMARY KEY" in col_def.upper()
                    nullable = "NOT NULL" not in col_def.upper() and not is_pk
                    pf = ParsedField(
                        table_name=table_name,
                        field_name=fname.upper(),
                        data_type=ftype.upper(),
                        nullable=nullable,
                        is_pk=is_pk,
                        raw_line=col_def[:80],
                    )
                    table.fields.append(pf)

            result.tables.append(table)

        if not result.tables:
            result.warnings.append("No CREATE TABLE statements found — check the file format")

        return result

    # ── JSON parser ────────────────────────────────────────────
    def _parse_json(self, text: str) -> SchemaParseResult:
        """
        Handles various JSON schema formats:
        - Array of table objects: [{name, fields: [{field, type, ...}]}]
        - Dict keyed by table name: {TABLE: [{field, type}]}
        - OpenAPI-style schemas
        """
        result = SchemaParseResult()
        try:
            data = json.loads(text)
        except json.JSONDecodeError as e:
            result.warnings.append(f"JSON parse error: {e}")
            return result

        if isinstance(data, list):
            # Array of table objects
            for item in data:
                if not isinstance(item, dict):
                    continue
                tname = str(item.get("name") or item.get("table") or item.get("TABLE_NAME", "UNKNOWN")).upper()
                table = ParsedTable(name=tname, description=item.get("description", ""))
                for col in item.get("fields", item.get("columns", [])):
                    if isinstance(col, dict):
                        pf = ParsedField(
                            table_name=tname,
                            field_name=str(col.get("field") or col.get("name") or col.get("COLUMN_NAME", "")).upper(),
                            data_type=str(col.get("type") or col.get("data_type") or col.get("DATA_TYPE", "")),
                            description=col.get("description") or col.get("remarks") or "",
                        )
                        table.fields.append(pf)
                result.tables.append(table)

        elif isinstance(data, dict):
            for tname, cols in data.items():
                if not isinstance(cols, list):
                    continue
                table = ParsedTable(name=tname.upper())
                for col in cols:
                    if isinstance(col, dict):
                        fname = str(col.get("f") or col.get("field") or col.get("name") or "").upper()
                        ftype = str(col.get("t") or col.get("type") or "")
                        fdesc = col.get("r") or col.get("description") or ""
                        if fname:
                            table.fields.append(ParsedField(
                                table_name=tname.upper(),
                                field_name=fname,
                                data_type=ftype,
                                description=str(fdesc),
                            ))
                result.tables.append(table)

        if not result.tables:
            result.warnings.append("No tables parsed from JSON — check structure")
        return result

    # ── XLSX parser ────────────────────────────────────────────
    def _parse_xlsx(self, content: bytes) -> SchemaParseResult:
        """
        Parses XLSX data dictionaries.
        Expects: each sheet = one table, or a single sheet with Table/Field/Type/Description columns.
        """
        result = SchemaParseResult()
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        except Exception as e:
            result.warnings.append(f"XLSX open error: {e}")
            return result

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue

            header = [str(c).upper().strip() if c else "" for c in rows[0]]

            # Check for multi-table layout (TABLE, FIELD, TYPE, DESCRIPTION columns)
            if "TABLE" in header or "TABLE_NAME" in header:
                ti = next((i for i, h in enumerate(header) if "TABLE" in h), 0)
                fi = next((i for i, h in enumerate(header) if "FIELD" in h or "COLUMN" in h), 1)
                ty = next((i for i, h in enumerate(header) if "TYPE" in h), 2)
                di = next((i for i, h in enumerate(header) if "DESC" in h or "REMARK" in h), -1)
                tables: dict[str, ParsedTable] = {}
                for row in rows[1:]:
                    if not row or not row[ti]:
                        continue
                    tname = str(row[ti]).upper()
                    fname = str(row[fi]).upper() if row[fi] else ""
                    ftype = str(row[ty]) if row[ty] else ""
                    fdesc = str(row[di]) if di >= 0 and row[di] else ""
                    if tname not in tables:
                        tables[tname] = ParsedTable(name=tname)
                    if fname:
                        tables[tname].fields.append(
                            ParsedField(table_name=tname, field_name=fname, data_type=ftype, description=fdesc)
                        )
                result.tables.extend(tables.values())
            else:
                # Each sheet is a table — columns are fields
                table = ParsedTable(name=sheet_name.upper())
                fi = next((i for i, h in enumerate(header) if "FIELD" in h or "NAME" in h or "COLUMN" in h), 0)
                ty = next((i for i, h in enumerate(header) if "TYPE" in h), 1)
                di = next((i for i, h in enumerate(header) if "DESC" in h or "REMARK" in h), -1)
                for row in rows[1:]:
                    if not row or not row[fi]:
                        continue
                    fname = str(row[fi]).upper()
                    ftype = str(row[ty]) if row[ty] else ""
                    fdesc = str(row[di]) if di >= 0 and row[di] else ""
                    table.fields.append(
                        ParsedField(table_name=sheet_name.upper(), field_name=fname, data_type=ftype, description=fdesc)
                    )
                result.tables.append(table)

        if not result.tables:
            result.warnings.append("No tables found in XLSX")
        return result

    # ── CSV parser ─────────────────────────────────────────────
    def _parse_csv(self, text: str) -> SchemaParseResult:
        """Simple CSV: TABLE,FIELD,TYPE,DESCRIPTION"""
        import csv
        result = SchemaParseResult()
        reader = csv.DictReader(io.StringIO(text))
        tables: dict[str, ParsedTable] = {}
        for row in reader:
            # Try various column name patterns
            tname = (row.get("TABLE") or row.get("TABLE_NAME") or row.get("table") or "UNKNOWN").upper()
            fname = (row.get("FIELD") or row.get("COLUMN_NAME") or row.get("field") or "").upper()
            ftype = row.get("TYPE") or row.get("DATA_TYPE") or row.get("type") or ""
            fdesc = row.get("DESCRIPTION") or row.get("REMARKS") or row.get("description") or ""
            if tname not in tables:
                tables[tname] = ParsedTable(name=tname)
            if fname:
                tables[tname].fields.append(
                    ParsedField(table_name=tname, field_name=fname, data_type=str(ftype), description=str(fdesc))
                )
        result.tables.extend(tables.values())
        return result

    # ── TXT parser ─────────────────────────────────────────────
    def _parse_txt(self, text: str) -> SchemaParseResult:
        """
        Multi-strategy text parser. Tries in order:
          1. SQL DDL (CREATE TABLE)
          2. Data dictionary format (table headers + field/type rows)
          3. Embedded JSON
        Handles Relius schema PDFs, data dictionaries, and exported DDL.
        """
        # Strategy 0: Frp record-layout (field codes like AA005 + Description)
        if self._looks_like_frp_layout(text):
            r = self._parse_frp_layout(text)
            if r.tables:
                return r

        # Strategy 1: SQL DDL
        if re.search(r'CREATE\s+TABLE', text, re.I):
            return self._parse_sql(text)

        # Strategy 2: Data dictionary format
        # Detects: table names as standalone all-caps words, followed by
        # field rows containing known SQL data types
        result = self._parse_data_dictionary(text)
        if result.tables:
            return result

        # Strategy 3: Embedded JSON
        json_match = re.search(r'\[.*?\]|\{.*?\}', text, re.DOTALL)
        if json_match:
            try:
                r = self._parse_json(json_match.group(0))
                if r.tables:
                    return r
            except Exception:
                pass

        result = SchemaParseResult(
            warnings=["Could not detect schema structure in this file. "
                      "For best results, export as SQL DDL or XLSX data dictionary."]
        )
        return result

    # Data type keywords that identify a field definition line
    _TYPE_PATTERN = re.compile(
        r"\b(CHAR|VARCHAR|VARCHAR2|NVARCHAR|NCHAR|TEXT|CLOB|BLOB|"
        r"NUMBER|NUMERIC|DECIMAL|INTEGER|INT|SMALLINT|BIGINT|TINYINT|"
        r"FLOAT|DOUBLE|REAL|MONEY|"
        r"DATE|DATETIME|TIMESTAMP|TIME|"
        r"BOOLEAN|BOOL|BIT)\b",
        re.I,
    )

    # Table name: 3+ uppercase letters/digits/underscores, alone on a line
    # (optionally followed by a dash or description)
    _TABLE_NAME_PATTERN = re.compile(
        r"^[ \t]*([A-Z][A-Z0-9_]{2,})[ \t]*(?:[-–—].*)?$",
        re.M,
    )

    def _parse_data_dictionary(self, text: str) -> SchemaParseResult:
        """
        Parse data dictionary / schema report format.

        Handles output like:
          TABLENAME
          FIELDNAME   CHAR(10)   Description text
          FIELDNAME2  INTEGER    Description text

        Or with explicit headers:
          TABLENAME
          Field        Type       Description
          FIELDNAME    CHAR(10)   Description text
        """
        result = SchemaParseResult()

        # Split text into lines and clean up
        lines = [l.rstrip() for l in text.splitlines()]

        current_table: Optional[ParsedTable] = None
        skip_next_header = False    # skip "Field / Type / Description" header rows

        # Words that look like table names but are actually headers/labels
        NOISE_WORDS = {
            "FIELD", "COLUMN", "NAME", "TYPE", "NULL", "KEY", "DEFAULT",
            "DESCRIPTION", "REMARK", "REMARKS", "COMMENT", "INDEX", "TABLE",
            "YES", "NO", "PK", "FK", "UNIQUE", "CONSTRAINT", "PRIMARY",
            "FOREIGN", "REFERENCES", "NOT", "AND", "OR", "THE", "FOR",
            "DATA", "SIZE", "LENGTH", "FORMAT", "VALUE", "SECTION",
            "PAGE", "APPENDIX", "CHAPTER", "FIGURE", "VERSION",
        }

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # Skip obvious header/divider lines
            if re.match(r'^[-=_*]+$', stripped):
                continue

            # Skip lines that look like "Field   Type   Description" headers
            if re.match(r'(?:Field|Column|Name).*(?:Type|DataType)', stripped, re.I):
                skip_next_header = True
                continue

            # Check if this line is a field definition (contains a known type keyword)
            type_match = self._TYPE_PATTERN.search(stripped)

            if type_match and current_table is not None:
                # This looks like a field row
                # Extract: FIELDNAME  TYPE(size)  optional description
                field_match = re.match(
                    r'^[ \t]*([A-Z_$#][A-Z0-9_$#]*)[ \t]+'
                    r'((?:CHAR|VARCHAR|VARCHAR2|NVARCHAR|NCHAR|TEXT|CLOB|BLOB|'
                    r'NUMBER|NUMERIC|DECIMAL|INTEGER|INT|SMALLINT|BIGINT|TINYINT|'
                    r'FLOAT|DOUBLE|REAL|MONEY|DATE|DATETIME|TIMESTAMP|TIME|'
                    r'BOOLEAN|BOOL|BIT)[^\s,;]*(?:\([^)]*\))?)'
                    r'[ \t]*(.*)?$',
                    stripped, re.I
                )
                if field_match:
                    fname = field_match.group(1).upper()
                    ftype = field_match.group(2).upper()
                    # Everything after type: may contain NULL/PK flags then description
                    rest  = (field_match.group(3) or "").strip()
                    # Strip YES/NO/PK/FK markers from the start of description
                    desc  = re.sub(r'^(?:YES|NO|PK|FK|NULL|NOT\s+NULL|UNIQUE)?\s*', '', rest, flags=re.I).strip()

                    # Skip if field name is a noise word (false positive)
                    if fname not in NOISE_WORDS and len(fname) > 1:
                        is_pk = bool(re.search(r'\bPK\b|PRIMARY\s+KEY', stripped, re.I))
                        is_fk = bool(re.search(r'\bFK\b|FOREIGN\s+KEY|REFERENCES', stripped, re.I))
                        pf = ParsedField(
                            table_name=current_table.name,
                            field_name=fname,
                            data_type=ftype,
                            nullable="NOT NULL" not in stripped.upper() and not is_pk,
                            is_pk=is_pk,
                            is_fk=is_fk,
                            description=desc[:200],
                        )
                        current_table.fields.append(pf)
                    continue

            # Check if this line is a standalone table name
            tbl_match = self._TABLE_NAME_PATTERN.match(line)
            if tbl_match:
                candidate = tbl_match.group(1).upper()

                # Must not be a noise word
                if candidate in NOISE_WORDS:
                    continue

                # Must look like a real table name (not a single word that's a type keyword)
                if self._TYPE_PATTERN.match(candidate):
                    continue

                # Minimum 3 chars, at least partially alphabetic
                if len(candidate) < 3 or not re.search(r'[A-Z]', candidate):
                    continue

                # Save current table if it has fields
                if current_table is not None and current_table.fields:
                    result.tables.append(current_table)
                elif current_table is not None:
                    # Table with no fields — still add it (may parse fields later)
                    result.tables.append(current_table)

                current_table = ParsedTable(name=candidate)
                skip_next_header = False
                continue

        # Don't forget the last table
        if current_table is not None:
            result.tables.append(current_table)

        # Post-process: remove tables with no fields if there are tables with fields
        tables_with_fields = [t for t in result.tables if t.fields]
        if tables_with_fields:
            result.tables = result.tables  # keep all tables including empty ones
        else:
            # No fields found anywhere — parser didn't work well
            result.tables = []
            result.warnings.append(
                "Found table names but could not extract field definitions. "
                "The PDF layout may not be in a standard tabular format."
            )

        return result

    # ── PDF parser ─────────────────────────────────────────────
    def _parse_pdf(self, content: bytes) -> SchemaParseResult:
        """
        Extract text from PDF then parse.
        Detects Relius 'Database Layout (By Tables)' format automatically.
        """
        pages_text: list[str] = []
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            pages_text = [page.extract_text() or "" for page in reader.pages]
            full_text = "\n".join(pages_text)
        except Exception as e:
            full_text = ""
            logger.warning("schema.pdf.text_error", error=str(e))

        # Image-only / scanned PDF (little or no embedded text): OCR the pages.
        # A text-layer heading alone can exceed a tiny threshold while the body
        # lives in page images, so also OCR when any page is near-empty, and
        # keep whichever source recovers more text.
        ocr_warn = ""
        sparse_pages = sum(1 for t in pages_text if len(t.strip()) < 40)
        if len(full_text.strip()) < 200 or sparse_pages:
            ocr_text, ocr_warn = self._ocr_pdf_pages(content)
            if len(ocr_text.strip()) > len(full_text.strip()):
                full_text = ocr_text
            elif not full_text.strip():
                return SchemaParseResult(warnings=[ocr_warn or "PDF read error: no extractable text"])

        # Detect Relius 'Database Layout (By Tables)' format
        if "Database Layout" in full_text and "By Tables" in full_text:
            # Use the module-level whitelist of known Relius table names.
            # This is a frozenset embedded directly in this module — no runtime
            # import dependency, no fallback to None.
            result = self._parse_relius_layout(full_text, valid_tables=_get_valid_tables())
            if result.tables:
                return result

        # Fall back to generic text parser
        result = self._parse_txt(full_text)
        if not result.tables:
            result.warnings.append(
                "PDF extracted but no schema structure found. "
                "Try exporting as SQL DDL or XLSX."
            )
        return result

    def _parse_relius_layout(
        self,
        text: str,
        valid_tables: set | None = None,
    ) -> SchemaParseResult:
        """
        Parse Relius 'Database Layout (By Tables)' PDF format.

        Each line:  TYPE  [DESCRIPTION]  ROWNUM[TABLENAME]  FIELDNAME  [LENGTH]
        - TYPE keyword (possibly "V ARCHAR2" with injected space from pypdf)
        - TABLENAME only appears on the first occurrence of each table per page
        - ROWNUM is the field sequence within that table
        - When description wraps to next line, that line is: ROWNUM FIELDNAME [LEN]

        valid_tables: whitelist of known table names. When provided, a table-name
        candidate is only accepted if it appears in this set. This prevents
        description words from being misidentified as table names.
        """
        result   = SchemaParseResult()
        tables: dict[str, ParsedTable] = {}
        current_table_name = ""

        # Noise: words that must never be treated as table or field names
        NAME_NOISE = {
            "DATE", "NUMBER", "VARCHAR", "VARCHAR2", "NVARCHAR", "CHAR",
            "INTEGER", "INT", "FLOAT", "CLOB", "BLOB", "TIMESTAMP",
            "DATETIME", "SMALLINT", "BIGINT", "NUMERIC", "DECIMAL",
            "BOOLEAN", "BIT", "TEXT", "NCHAR", "TIME",
            # Date format strings that appear in descriptions
            "YYYYMMDD", "MMDDYYYY", "HHMMSS", "YYYYMM", "MMYYYY",
            "YYYYMMDDHHMMSS",
            # Common English words that appear in descriptions
            "TRUE", "FALSE", "NULL", "NONE", "YES", "NO",
            "FORMAT", "VALUE", "FIELD", "TABLE", "COLUMN", "TYPE",
            "INTEREST", "SPECIFIED", "REACT", "NOTE",
        }

        def valid_name(name: str, min_len: int = 2) -> bool:
            """Return True if name is a plausible Relius table or field name."""
            if not name: return False
            if name in NAME_NOISE: return False
            if len(name) < min_len or len(name) > 50: return False
            if not re.match(r'^[A-Z][A-Z0-9_]+$', name): return False
            # Reject date-format patterns
            if re.search(r'YYYY|MMDD|HHMM|HHSS', name): return False
            return True

        # TYPE keyword at start (handle "V ARCHAR2" space variant from pypdf)
        TYPE_PAT = re.compile(
            r"^(V\s*ARCHAR2?|N\s*VARCHAR2?|NUMBER|DATE(?:/TIME)?|CHAR|"
            r"INTEGER|INT|FLOAT|CLOB|BLOB|TIMESTAMP|DATETIME|"
            r"SMALLINT|BIGINT|NUMERIC|DECIMAL|BOOLEAN|BIT)(\d*)",
            re.I,
        )
        # End-of-line: ROWNUM [TABLENAME] FIELDNAME [LENGTH]
        TAIL_PAT = re.compile(
            r"(\d+)"
            r"([A-Z][A-Z0-9_]*(?:_[A-Z0-9]+)*)?"
            r"\s+([A-Z_$#][A-Z0-9_$#]{0,49})"
            r"(?:\s+([\d(),]+))?\s*$"
        )
        # Skip lines: page headers, column headers, blank
        SKIP_RE = re.compile(
            r"^(?:[\d/]{4,}Relius|Database Layout|\(By Tables\)|"
            r"Table Name\s+Field|\s*$)",
            re.I,
        )

        for line in text.splitlines():
            # Normalise pypdf VA-ligature splits before any parsing
            # e.g. "COMBOV ALUESDET" -> "COMBOVALUESDET", "V ALRULENUM" -> "VALRULENUM"
            stripped = _normalise_pypdf_va(line.strip())
            if not stripped or SKIP_RE.match(stripped):
                continue

            parsed_table = ""
            parsed_field = ""
            parsed_type  = ""
            parsed_desc  = ""
            parsed_len   = ""

            type_m = TYPE_PAT.match(stripped)
            if type_m:
                dtype      = type_m.group(1).replace(" ", "").upper()
                inline_seq = type_m.group(2)   # digits glued to type: "NUMBER5"
                rest       = stripped[type_m.end():].strip()

                if inline_seq:
                    # "NUMBER5 FIELDNAME [LEN]" or "NUMBER1TABLENAME FIELDNAME [LEN]"
                    # Split rest into tokens, ignore length tokens
                    tokens = [t for t in rest.split() if t]
                    word_tokens = [t for t in tokens if not re.match(r'^[\d(),]+$', t)]
                    len_token   = tokens[-1] if tokens and re.match(r'^[\d(),]+$', tokens[-1]) else ""

                    if len(word_tokens) >= 2:
                        # Could be TABLE FIELD — validate both
                        cand_table = word_tokens[0].upper()
                        cand_field = word_tokens[1].upper()
                        ct_ok = valid_name(cand_table, 2) and (valid_tables is None or cand_table in valid_tables)
                        if ct_ok and valid_name(cand_field, 1):
                            parsed_table = cand_table
                            parsed_field = cand_field
                            parsed_len   = len_token
                            parsed_type  = dtype
                        elif valid_name(cand_field, 1):
                            # First token invalid as table — treat both as field
                            parsed_field = cand_field
                            parsed_len   = len_token
                            parsed_type  = dtype
                    elif len(word_tokens) == 1:
                        cand_field = word_tokens[0].upper()
                        if valid_name(cand_field, 1):
                            parsed_field = cand_field
                            parsed_len   = len_token
                            parsed_type  = dtype
                else:
                    # Normal: TYPE [desc] ROWNUM[TABLE] FIELD [LEN]
                    tail_m = TAIL_PAT.search(rest)
                    if tail_m:
                        cand_table = (tail_m.group(2) or "").upper()
                        cand_field = tail_m.group(3).upper()
                        parsed_len = tail_m.group(4) or ""
                        parsed_desc = rest[:tail_m.start()].strip()
                        parsed_type = dtype

                        if valid_name(cand_table, 2) and (valid_tables is None or cand_table in valid_tables):
                            parsed_table = cand_table
                        if valid_name(cand_field, 1):
                            parsed_field = cand_field
                        else:
                            continue   # field name invalid — skip line
                    else:
                        # TAIL_PAT failed — pypdf may have split the table or field name
                        # e.g. "VALRULENUM" → "V ALRULENUM", "CMSACTVASSOC" → "CMSA CTVASSOC"
                        # Try to recover by joining split tokens using the whitelist.
                        ft, ff, fl = _try_tail_with_split_fix(rest, valid_tables)
                        if ft or ff:
                            parsed_table = ft or ""
                            parsed_field = ff or ""
                            parsed_len   = fl
                            parsed_type  = dtype

            else:
                # No type prefix — could be a continuation line:
                # "2 FIELDNAME [LEN]"  or  "2 TABLENAME FIELDNAME [LEN]"
                # Must start with digits
                cont_m = re.match(r'^(\d+)\s*(.*)', stripped)
                if not cont_m:
                    continue

                rest_cont  = cont_m.group(2).strip()
                tokens     = [t for t in rest_cont.split() if t]
                word_tokens = [t for t in tokens if not re.match(r'^[\d(),]+$', t)]
                len_token   = tokens[-1] if tokens and re.match(r'^[\d(),]+$', tokens[-1]) else ""

                if len(word_tokens) >= 2:
                    # Two words: could be TABLE FIELD (new table starts in continuation)
                    cand_table = word_tokens[0].upper()
                    cand_field = word_tokens[1].upper()
                    ct_ok = valid_name(cand_table, 2) and (valid_tables is None or cand_table in valid_tables)
                    if ct_ok and valid_name(cand_field, 1):
                        parsed_table = cand_table
                        parsed_field = cand_field
                        parsed_len   = len_token
                    elif valid_name(cand_field, 1):
                        parsed_field = cand_field
                        parsed_len   = len_token
                elif len(word_tokens) == 1:
                    cand_field = word_tokens[0].upper()
                    if valid_name(cand_field, 1):
                        parsed_field = cand_field
                        parsed_len   = len_token

            # Must have a field name to proceed
            if not parsed_field:
                continue

            # Update table tracker
            if parsed_table:
                current_table_name = parsed_table
            if not current_table_name:
                continue

            # Create table entry if new
            if current_table_name not in tables:
                tables[current_table_name] = ParsedTable(name=current_table_name)

            # Add field
            tables[current_table_name].fields.append(ParsedField(
                table_name  = current_table_name,
                field_name  = parsed_field,
                data_type   = parsed_type,
                description = parsed_desc,
            ))

        result.tables = list(tables.values())
        result.raw_format = "pdf-relius"

        if not result.tables:
            result.warnings.append(
                "Relius layout detected but no tables extracted. "
                "Check the PDF is not a scanned image."
            )

        return result


    # ── DOCX parser ────────────────────────────────────────────
    def _parse_docx(self, content: bytes) -> SchemaParseResult:
        """
        Parse a DOCX. Two shapes are handled automatically:
          - a real data dictionary (paragraphs / tables of text), and
          - a doc whose schema info lives in *screenshots* (embedded images),
            which are OCR'd via tesseract and folded into the text.
        """
        warnings: list[str] = []
        parts: list[str] = []
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            parts.extend(p.text for p in doc.paragraphs)
            for tbl in doc.tables:
                for row in tbl.rows:
                    parts.append("\t".join(c.text for c in row.cells))
        except Exception as e:
            return SchemaParseResult(warnings=[f"DOCX extraction error: {e}"])

        text_only = "\n".join(p for p in parts if p.strip())

        # OCR any embedded screenshots and append the recognised text.
        ocr_text, ocr_warn = self._ocr_docx_images(content)
        if ocr_warn:
            warnings.append(ocr_warn)

        combined = "\n".join(t for t in (text_only, ocr_text) if t.strip())
        result = self._parse_txt(combined) if combined.strip() else SchemaParseResult()
        result.warnings.extend(warnings)
        if not result.tables:
            result.warnings.append(
                "DOCX extracted but no structured schema found. "
                "If the tables are screenshots, ensure tesseract OCR is installed; "
                "otherwise export as a tabular XLSX/CSV for best results."
            )
        return result

    def _ocr_docx_images(self, content: bytes) -> tuple[str, str]:
        """
        OCR every image embedded in a .docx (they live under word/media/).
        Returns (recognised_text, warning). Degrades gracefully when the
        OCR stack (pytesseract + Pillow + the tesseract binary) is absent.
        """
        import zipfile
        images: list[bytes] = []
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as z:
                for name in z.namelist():
                    if name.lower().startswith("word/media/") and name.lower().endswith(
                        (".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".gif")
                    ):
                        images.append(z.read(name))
        except Exception as e:
            return "", f"Could not read embedded images: {e}"

        if not images:
            return "", ""

        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            return "", (
                f"Found {len(images)} screenshot(s) but OCR is not available — "
                "install pytesseract + Pillow and the tesseract binary "
                "(`brew install tesseract`) to read schema info from images."
            )

        def _prep(im: "Image.Image") -> "Image.Image":
            # Grayscale + upscale small screenshots — both materially improve
            # tesseract accuracy on UI/table captures.
            im = im.convert("L")
            w, h = im.size
            if w and w < 1600:
                scale = min(3.0, 1600 / w)
                im = im.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
            return im

        chunks: list[str] = []
        failures = 0
        for raw in images:
            try:
                im = _prep(Image.open(io.BytesIO(raw)))
                # psm 6 = "assume a single uniform block of text" — best fit for
                # tabular record/field screenshots.
                txt = pytesseract.image_to_string(im, config="--psm 6")
                if txt.strip():
                    chunks.append(txt)
            except Exception:
                failures += 1
        warn = ""
        if failures:
            warn = (
                f"OCR failed on {failures}/{len(images)} image(s) — "
                "is the tesseract binary installed and on PATH?"
            )
        logger.info("schema.docx.ocr", images=len(images), recognised=len(chunks), failures=failures)
        return "\n".join(chunks), warn

    # ── COBOL parser — routes by content (copybook vs program) ─
    # ── Frp record-layout parser ──────────────────────────────
    # Frp data-dictionary docs list each record's data elements as
    #   <CODE> <Name>            (code = 2 letters + 3 digits, e.g. AA005)
    #     Description
    #     <free text>
    #     [Legal Values]
    #     <value> <label> ...
    # Fields are grouped into records by the code's 2-letter prefix, which is
    # robust to garbled/absent section headers (important for OCR'd screenshots).
    _FRP_CODE = re.compile(r'^([A-Z]{2}\d{3}[A-Z0-9]{0,2})\s+([A-Za-z][^\n]{0,79})$')
    _FRP_PREFIX_NAME = {
        "AA": "Alternate Address", "AI": "Associated Individuals", "AR": "Auto Rebalance",
        "AM": "Annuity Master", "AP": "Annuity Pointer", "BT": "Base Text",
        "CD": "Check Detail", "CK": "Check Header", "EQ": "Equity Wash",
        "FC": "Fund Control", "PF": "Participant Fund", "PH": "Participant Header",
        "PL": "Plan", "PE": "Person", "DS": "Division", "SA": "Share Account",
        "BR": "History Base", "DB": "Disbursement", "CM": "Compliance",
        "SL": "Salary", "FR": "Forecast",
    }

    def _looks_like_frp_layout(self, text: str) -> bool:
        codes = re.findall(r'\b[A-Z]{2}\d{3}\b', text)
        return len(set(codes)) >= 3 and ("Description" in text or "Legal Value" in text)

    # Sentence starters that indicate a line is prose, not a record header.
    _FRP_NOISE_START = {
        "this", "the", "if", "note", "refer", "a", "an", "for", "when",
        "overridden", "indicates", "contains", "positions", "valid",
    }

    def _is_frp_header(self, line: str) -> bool:
        """A record section header like 'ADDRESS RECORD' / 'Associated Individuals Record'."""
        if not re.search(r"\brecords?\b", line, re.I):
            return False
        words = line.split()
        if not (1 < len(words) <= 6):
            return False
        if words[0].lower() in self._FRP_NOISE_START:
            return False
        return line[:1].isupper()

    def _parse_frp_layout(self, text: str) -> SchemaParseResult:
        result = SchemaParseResult(raw_format="frp-layout")
        records: dict[str, ParsedTable] = {}
        order: list[str] = []
        state = {"field": None, "mode": None, "legal": []}

        def flush_legal():
            fld = state["field"]
            if fld is not None and state["legal"]:
                fld.description = (fld.description + "  Legal values: " + "; ".join(state["legal"])).strip()
            state["legal"] = []

        for raw in text.splitlines():
            line = raw.strip()
            if not line:
                continue
            m = self._FRP_CODE.match(line)
            # A record header ends the current field's description/legal capture.
            if not m and self._is_frp_header(line):
                flush_legal()
                state["field"] = None
                state["mode"] = None
                continue
            if m:
                flush_legal()
                code, name = m.group(1), m.group(2).strip()
                prefix = code[:2]
                if prefix not in records:
                    records[prefix] = ParsedTable(name=self._FRP_PREFIX_NAME.get(prefix, prefix))
                    order.append(prefix)
                fld = ParsedField(
                    table_name=records[prefix].name,
                    field_name=f"{code} {name}".strip(),
                    description="",
                )
                records[prefix].fields.append(fld)
                state["field"] = fld
                state["mode"] = None
                continue
            low = line.lower()
            if low == "description":
                flush_legal(); state["mode"] = "desc"; continue
            if low.startswith("legal value"):
                flush_legal(); state["mode"] = "legal"; continue
            fld = state["field"]
            if state["mode"] == "desc" and fld is not None:
                fld.description = (fld.description + " " + line).strip()[:2000]
            elif state["mode"] == "legal" and fld is not None:
                state["legal"].append(re.sub(r"\s+", " ", line))
        flush_legal()

        result.tables = [records[k] for k in order if records[k].fields]
        if not result.tables:
            result.warnings.append("Frp layout detected but no records/fields could be extracted.")
        return result

    def _ocr_pdf_pages(self, content: bytes) -> tuple[str, str]:
        """
        Render each PDF page to an image and OCR it — for image-only / scanned
        PDFs where pypdf finds no text. Degrades gracefully if the OCR stack or
        PyMuPDF is unavailable.
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            return "", ("PDF appears image-only but PyMuPDF is not installed — "
                        "run `pip install pymupdf` to OCR scanned PDFs.")
        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            return "", ("PDF appears image-only but OCR is not available — install "
                        "pytesseract + Pillow and the tesseract binary (`brew install tesseract`).")

        chunks: list[str] = []
        failures = 0
        try:
            doc = fitz.open(stream=content, filetype="pdf")
        except Exception as e:
            return "", f"Could not open PDF for OCR: {e}"
        for page in doc:
            try:
                # ~300 DPI render for legible OCR of small screenshot text
                pix = page.get_pixmap(matrix=fitz.Matrix(300 / 72, 300 / 72))
                img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("L")
                txt = pytesseract.image_to_string(img, config="--psm 6")
                if txt.strip():
                    chunks.append(txt)
            except Exception:
                failures += 1
        warn = f"OCR failed on {failures} page(s)." if failures else ""
        logger.info("schema.pdf.ocr", pages=len(chunks), failures=failures)
        return "\n".join(chunks), warn

    def _parse_cobol(self, text: str) -> SchemaParseResult:
        """
        Decide which COBOL parser to use based on the file's content:
          - a copybook / record layout → data-definition parser, and
          - a program holding mapping logic (PROCEDURE DIVISION / MOVE …
            TO …) → logic parser that also reads its data definitions.
        """
        upper = text.upper()
        is_program = (
            re.search(r'\bPROCEDURE\s+DIVISION\b', upper) is not None
            or len(re.findall(r'\bMOVE\b[^.\n]+?\bTO\b', upper)) >= 2
        )
        return self._parse_cobol_program(text) if is_program else self._parse_cobol_copybook(text)

    # ── COBOL copybook parser (data definitions) ───────────────
    def _parse_cobol_copybook(self, text: str) -> SchemaParseResult:
        """
        Parse a COBOL copybook / record layout into tables + fields.

        Each 01-level item becomes a table (record); each elementary item
        (one carrying a PIC clause) becomes a field whose PIC clause is
        mapped to an approximate SQL data type. Level 88 (condition names),
        66 (RENAMES), group items (no PIC) and FILLER are skipped.

        Handles both fixed-format copybooks (6-digit sequence area + an
        indicator column) and free-format ones.
        """
        result = SchemaParseResult()
        result.raw_format = "cobol"

        tables: list[ParsedTable] = []
        current: Optional[ParsedTable] = None

        ENTRY_RE = re.compile(r'^(\d{1,2})\s+([A-Z0-9][A-Z0-9-]*)\b(.*)$', re.I)
        PIC_RE   = re.compile(r'\bPIC(?:TURE)?\s+(?:IS\s+)?([^\s.]+)', re.I)

        for raw in text.splitlines():
            line = raw
            # Fixed-format comment / continuation indicator lives in column 7
            if len(line) > 6 and line[6] in ("*", "/"):
                continue
            # Strip a numeric sequence-number area (cols 1-6) + indicator column
            if re.match(r'^\d{6}', line):
                line = line[7:] if len(line) > 6 else ""

            stmt = line.strip()
            if not stmt or stmt.startswith("*"):
                continue

            # One data-description entry ends at its terminating period
            stmt = stmt.split(".", 1)[0].strip()
            m = ENTRY_RE.match(stmt)
            if not m:
                continue

            level = int(m.group(1))
            name  = m.group(2).upper()
            rest  = m.group(3) or ""

            if level in (88, 66):          # condition-name / RENAMES — not fields
                continue

            if level == 1:                 # 01 → new record == table
                if current is not None:
                    tables.append(current)
                current = ParsedTable(name=name.replace("-", "_"))
                continue

            if name == "FILLER":
                continue

            if current is None:            # 77-level or fragment without an 01
                current = ParsedTable(name="COPYBOOK_RECORD")

            pic_m = PIC_RE.search(rest)
            if not pic_m:                  # group item — only PIC-bearing items are fields
                continue

            current.fields.append(ParsedField(
                table_name = current.name,
                field_name = name.replace("-", "_"),
                data_type  = self._cobol_pic_to_type(pic_m.group(1)),
                raw_line   = stmt[:80],
            ))

        if current is not None:
            tables.append(current)

        result.tables = tables
        if not result.tables:
            result.warnings.append(
                "No COBOL record (01-level) definitions found. "
                "Ensure the file is a copybook with PIC clauses."
            )
        return result

    # ── COBOL program parser (Relius → Frp mapping logic) ─────
    def _parse_cobol_program(self, text: str) -> SchemaParseResult:
        """
        Parse a COBOL program that holds Relius → Frp mapping logic.

        Strategy:
          1. Extract any record/field definitions present (WORKING-STORAGE /
             LINKAGE 01-levels) using the copybook parser.
          2. Scan PROCEDURE DIVISION `MOVE <src> TO <tgt>` statements. The
             receiving (target) items are Frp fields; the sending item is
             recorded as a mapping hint ("← SRC") on each target field.
          3. Targets that have no data definition are synthesised into a
             record (the OF/IN qualifier, else FRP_MAPPING) so the file
             still yields a usable Frp field list.
        """
        result = self._parse_cobol_copybook(text)
        result.raw_format = "cobol-program"

        # Index existing tables/fields for enrichment / dedupe
        existing_tables: dict[str, ParsedTable] = {t.name: t for t in result.tables}
        by_qualified: dict[tuple, ParsedField] = {}
        by_name: dict[str, ParsedField] = {}
        for t in result.tables:
            for f in t.fields:
                by_qualified[(t.name, f.field_name)] = f
                by_name.setdefault(f.field_name, f)
        synth: dict[str, ParsedTable] = {}

        IDENT = r"[A-Za-z][A-Za-z0-9-]*"
        MOVE_RE = re.compile(rf"\bMOVE\s+(.+?)\s+TO\s+([^.]+)", re.I)
        TGT_RE  = re.compile(rf"({IDENT})(?:\s+(?:OF|IN)\s+({IDENT}))?", re.I)

        def norm(s: str) -> str:
            return s.upper().replace("-", "_")

        def source_label(raw: str) -> str:
            raw = raw.strip()
            m = re.match(rf"({IDENT})(?:\s+(?:OF|IN)\s+({IDENT}))?", raw, re.I)
            if m:
                return norm(m.group(2) + "." + m.group(1)) if m.group(2) else norm(m.group(1))
            return raw[:40]  # literal / figurative constant (e.g. SPACES, 0, 'C')

        mapping_count = 0
        for mv in MOVE_RE.finditer(text):
            src = source_label(mv.group(1))
            for tm in TGT_RE.finditer(mv.group(2)):
                tfield = norm(tm.group(1))
                trec   = norm(tm.group(2)) if tm.group(2) else None
                # COBOL verbs sometimes trail the receiver (e.g. ROUNDED) — ignore them
                if tfield in {"ROUNDED", "GIVING", "DELIMITED"}:
                    continue
                note = f"← {src}"
                target = (by_qualified.get((trec, tfield)) if trec else None) or by_name.get(tfield)
                if target is not None:
                    target.description = (
                        f"{target.description}; {note}".strip("; ") if target.description else note
                    )
                else:
                    rec = trec or "FRP_MAPPING"
                    # Attach to the existing record if the qualifier names one,
                    # else group into a synthesised mapping table.
                    tbl = existing_tables.get(rec) or synth.setdefault(rec, ParsedTable(name=rec))
                    if not any(f.field_name == tfield for f in tbl.fields):
                        tbl.fields.append(ParsedField(
                            table_name=rec, field_name=tfield, data_type="", description=note,
                        ))
                        by_qualified[(rec, tfield)] = tbl.fields[-1]
                        by_name.setdefault(tfield, tbl.fields[-1])
                mapping_count += 1

        result.tables.extend(synth.values())
        if mapping_count:
            result.warnings.append(
                f"Detected COBOL mapping logic — captured {mapping_count} MOVE mapping(s) "
                "as Frp field hints (source field shown as '← …' in the description)."
            )
        if not result.tables:
            result.warnings.append(
                "COBOL program detected but no record definitions or MOVE mappings "
                "could be extracted."
            )
        return result

    @staticmethod
    def _cobol_pic_to_type(pic: str) -> str:
        """Map a COBOL PICTURE clause to an approximate SQL data type."""
        p = pic.upper().strip().rstrip(".")
        # Expand repetition factors: X(10) -> XXXXXXXXXX, 9(5) -> 99999
        expanded = re.sub(r'(.)\((\d+)\)', lambda mm: mm.group(1) * int(mm.group(2)), p)
        if "X" in expanded:
            n = expanded.count("X")
            return f"CHAR({n})" if n else "CHAR"
        if "A" in expanded:
            n = expanded.count("A")
            return f"CHAR({n})" if n else "CHAR"
        if "9" in expanded:
            if "V" in expanded:            # implied decimal point
                intp, _, frac = expanded.partition("V")
                ip, fp = intp.count("9"), frac.count("9")
                return f"DECIMAL({ip + fp},{fp})"
            return f"NUMBER({expanded.count('9')})"
        return p or "UNKNOWN"


# Module-level singleton
extractor = SchemaExtractor()